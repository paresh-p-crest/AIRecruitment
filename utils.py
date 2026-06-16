"""Document parsing helpers for PDF, DOCX, and DOC resume uploads."""

import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import HTTPException, UploadFile, status

from doc_converter import extract_text_from_doc
from resume_parser import clean_resume_text

logger = logging.getLogger(__name__)

PDF_EXTRACTORS = ("pymupdf", "pdfplumber", "pypdf")

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


def _validate_extension(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Unsupported file type '{extension or 'unknown'}'. "
                f"Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
            ),
        )
    return extension


def _extract_pdf_with_pdfplumber(content: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
    return "\n\n".join(text_parts).strip()


def _extract_pdf_with_pymupdf(content: bytes) -> str:
    import fitz

    text_parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
    return "\n\n".join(text_parts).strip()


def _extract_pdf_with_pypdf(content: bytes) -> str:
    from pypdf import PdfReader

    text_parts: list[str] = []
    reader = PdfReader(io.BytesIO(content))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())
    return "\n\n".join(text_parts).strip()


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract plain text from a PDF using multiple parsers.

    Some resumes (e.g. certain Word-exported PDFs) fail with pdfplumber
    but work with PyMuPDF or pypdf — we try each in order.
    """
    # PyMuPDF first — many Word-exported resumes return empty text from pdfplumber.
    extractors = (
        ("pymupdf", _extract_pdf_with_pymupdf),
        ("pdfplumber", _extract_pdf_with_pdfplumber),
        ("pypdf", _extract_pdf_with_pypdf),
    )

    for name, extractor in extractors:
        try:
            text = extractor(content)
            if text:
                if name != "pymupdf":
                    logger.info("PDF text extracted using %s fallback", name)
                return text
        except Exception as exc:
            logger.warning("PDF extraction via %s failed: %s", name, exc)

    return ""


def _docx_hyperlink_targets(document: Document) -> list[str]:
    """Collect hyperlink URLs from DOCX (mailto: and visible link targets)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    targets: list[str] = []
    for paragraph in document.paragraphs:
        for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
            rel_id = hyperlink.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
            )
            if not rel_id:
                continue
            try:
                rel = paragraph.part.rels[rel_id]
            except KeyError:
                continue
            if rel.reltype == RT.HYPERLINK and rel.target_ref:
                targets.append(rel.target_ref)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                        rel_id = hyperlink.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        )
                        if not rel_id:
                            continue
                        try:
                            rel = paragraph.part.rels[rel_id]
                        except KeyError:
                            continue
                        if rel.reltype == RT.HYPERLINK and rel.target_ref:
                            targets.append(rel.target_ref)
    return targets


def extract_text_from_docx(content: bytes) -> str:
    """Extract plain text from a DOCX byte stream (paragraphs + tables + hyperlinks)."""
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    seen: set[str] = set()

    def _add_line(line: str) -> None:
        text = line.strip()
        if not text or text in seen:
            return
        seen.add(text)
        parts.append(text)

    for paragraph in document.paragraphs:
        _add_line(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                _add_line(" | ".join(cells))

    for target in _docx_hyperlink_targets(document):
        _add_line(target)

    return "\n".join(parts).strip()


def validate_resume_bytes(filename: str, content: bytes) -> str:
    """
    Validate upload bytes and return the normalized file extension.

    Raises ValueError for invalid or oversized files.
    """
    if not filename:
        raise ValueError("Uploaded file must include a filename.")

    base_name = Path(filename).name
    if base_name.startswith("~$"):
        raise ValueError(
            "This looks like a temporary Word lock file (~$...). "
            "Close the document in Word and upload the actual resume file."
        )

    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{extension or 'unknown'}'. "
            f"Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if not content:
        raise ValueError("Uploaded file is empty.")

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File exceeds maximum size of {MAX_FILE_SIZE_BYTES // (1024 * 1024)} MB."
        )

    return extension


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """
    Extract plain text from resume bytes.

    Raises ValueError when parsing fails or no text is found.
    """
    extension = validate_resume_bytes(filename, content)

    try:
        if extension == ".pdf":
            raw_text = extract_text_from_pdf(content)
        elif extension == ".docx":
            raw_text = extract_text_from_docx(content)
        else:
            raw_text, _method = extract_text_from_doc(content, filename)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Failed to parse {extension} file: {exc}") from exc

    if not raw_text:
        raise ValueError(
            "No readable text could be extracted from the uploaded document. "
            "If this is a scanned/image-only PDF, try exporting a text-based PDF or DOCX."
        )

    return clean_resume_text(raw_text)


async def extract_text_from_upload(file: UploadFile) -> tuple[str, str]:
    """
    Read an uploaded resume file and return (extension, extracted_text).

    Raises HTTPException for invalid file types, empty files, or unreadable content.
    """
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file must include a filename.",
        )

    content = await file.read()

    try:
        extension = validate_resume_bytes(file.filename, content)
        raw_text = extract_text_from_bytes(file.filename, content)
    except ValueError as exc:
        message = str(exc)
        if "Unsupported file type" in message:
            status_code = status.HTTP_400_BAD_REQUEST
        elif "exceeds maximum size" in message:
            status_code = status.HTTP_413_REQUEST_ENTITY_TOO_LARGE
        elif ".doc file" in message or "Could not extract text from .doc" in message:
            status_code = status.HTTP_503_SERVICE_UNAVAILABLE
        else:
            status_code = status.HTTP_422_UNPROCESSABLE_ENTITY
        raise HTTPException(status_code=status_code, detail=message) from exc

    return extension, raw_text
